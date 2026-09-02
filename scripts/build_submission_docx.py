from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "submission" / "unlimit-ai-assessment-proposal.md"
OUTPUT = ROOT / "docs" / "submission" / "unlimit-ai-assessment-proposal.docx"


def set_cell_shading(cell, fill: str) -> None:  # type: ignore[no-untyped-def]
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_inline(paragraph, text: str) -> None:  # type: ignore[no-untyped-def]
    parts = re.split(r"(\*\*.*?\*\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        else:
            paragraph.add_run(part)


def configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.02

    for name, size, color in (
        ("Heading 1", 15, "102A43"),
        ("Heading 2", 11.5, "126E82"),
        ("Heading 3", 10, "102A43"),
    ):
        style = document.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(7)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "UNLIMIT AI ASSESSMENT  |  SENTINEL"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(7)
    header.runs[0].font.color.rgb = RGBColor.from_string("627D98")

    footer = section.footer.paragraphs[0]
    footer.text = "Working proposal — 2 September 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.name = "Arial"
    footer.runs[0].font.size = Pt(7)
    footer.runs[0].font.color.rgb = RGBColor.from_string("829AB1")


def add_table(document: Document, rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, value in enumerate(rows[0]):
        cell = table.rows[0].cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, "D9EAF0")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(7.5)
    for row in rows[1:]:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(1)
                for run in paragraph.runs:
                    run.font.size = Pt(7.2)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def build() -> None:
    document = Document()
    configure(document)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 0
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        paragraph = document.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))
        paragraph_lines.clear()

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            flush_paragraph()
            if in_code:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.left_indent = Cm(0.35)
                paragraph.paragraph_format.space_before = Pt(2)
                paragraph.paragraph_format.space_after = Pt(5)
                set_cell = OxmlElement("w:shd")
                set_cell.set(qn("w:fill"), "F0F4F8")
                paragraph._p.get_or_add_pPr().append(set_cell)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
                run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
                run.font.size = Pt(7.2)
                code_lines.clear()
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: list[list[str]] = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                values = [value.strip() for value in lines[index].strip().strip("|").split("|")]
                if not all(set(value) <= {"-", ":"} for value in values):
                    table_lines.append(values)
                index += 1
            add_table(document, table_lines)
            continue
        if not stripped:
            flush_paragraph()
            index += 1
            continue
        if stripped.startswith("# "):
            flush_paragraph()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(stripped[2:])
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(20)
            run.font.color.rgb = RGBColor.from_string("102A43")
        elif stripped.startswith("## "):
            flush_paragraph()
            document.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            flush_paragraph()
            document.add_heading(stripped[4:], level=3)
        elif re.match(r"^\d+\. ", stripped):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Number")
            add_inline(paragraph, re.sub(r"^\d+\. ", "", stripped))
        elif stripped.startswith("- "):
            flush_paragraph()
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, stripped[2:])
        elif stripped.endswith("  "):
            paragraph_lines.append(stripped[:-2])
            flush_paragraph()
        else:
            paragraph_lines.append(stripped)
        index += 1
    flush_paragraph()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
